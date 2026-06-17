#!/usr/bin/env python3
"""Convert PLAN-COMERCIALIZACION.md to a formatted Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- Helper functions ---
def add_title(text, level=1):
    if level == 0:
        p = doc.add_heading(text, level=0)
    else:
        p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
    return p

def add_orange_heading(text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
    run.font.name = 'Calibri'
    return p

def add_body(text):
    p = doc.add_paragraph(text)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(11)
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Orange background
        shading = cell._element.get_or_add_tcPr()
        shade_elem = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): 'C2410C',
            qn('w:val'): 'clear'
        })
        shading.append(shade_elem)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing
    return table

# ============================================================
# DOCUMENT CONTENT
# ============================================================

# --- Cover / Title ---
doc.add_paragraph()  # spacer
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('PLAN DE COMERCIALIZACIÓN')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)
run.font.name = 'Calibri'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('QuickTable — Menú digital inteligente para restaurantes')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Análisis de costos, capacidad y rentabilidad para SaaS dirigido a restaurantes en Perú')
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('Junio 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# --- SECTION 1: VERSIÓN BÁSICA ---
add_orange_heading('VERSIÓN 1: BÁSICA (Solo Web — PWA)', 18)
add_body('Plataforma web progresiva (PWA) responsive que funciona desde el navegador en cualquier dispositivo (celular, tablet, PC). No necesita instalación desde App Store. Todo el backend corre en un solo VPS.')

add_orange_heading('Funcionalidades incluidas', 14)
features_basic = [
    'Menú digital con QR por mesa (cliente escanea y pide)',
    'Panel empresa: Dashboard KPIs, Cocina, Chat con IA',
    'Inventario completo: Stock, Kardex, Recetas, Mermas, Kanban',
    'Facturación electrónica Sunat vía API (Nubefact / Facturador.com)',
    'Pasarela de pagos: Mercado Pago + Culqi',
    'Multi-sede con jerarquía, roles y permisos (Superadmin → Admin de sede → Mozo → Cocina → Caja → Almacén)',
    'Chat con IA (OpenAI GPT-4o-mini)',
    'App responsive PWA (no necesita Google Play)',
]
for f in features_basic:
    add_bullet(f)

add_orange_heading('Infraestructura', 14)
infra_basic = [
    '1 VPS DigitalOcean / AWS Lightsail (8GB RAM, 4 vCPU, 160GB SSD)',
    'PostgreSQL en el mismo servidor',
    'SSL + dominio .pe',
]
for i in infra_basic:
    add_bullet(i)

add_orange_heading('Costos', 14)
cost_headers = ['Concepto', 'Costo único (S/)', 'Costo mensual (S/)']
cost_rows = [
    ['Desarrollo backend (FastAPI + PostgreSQL + WebSockets)', '8,000 – 12,000', '—'],
    ['Frontend web (React/Svelte + PWA responsive)', '5,000 – 8,000', '—'],
    ['Facturación Sunat (integración Nubefact/Facturador.com)', '2,000 – 3,500', '—'],
    ['Pasarela de pagos (Mercado Pago + Culqi)', '1,500 – 2,500', '—'],
    ['Multi-sede + jerarquía + roles + permisos', '2,000 – 3,500', '—'],
    ['IA Chat (GPT-4o-mini para 50 clientes)', '—', '1,200 – 1,800'],
    ['VPS (8GB RAM, 4 vCPU)', '—', '150 – 250'],
    ['Dominio .pe + SSL (1 año)', '120', '10'],
    ['Logo + identidad visual', '500', '—'],
    ['Landing page + documentación', '1,000 – 2,000', '—'],
    ['Soporte técnico (dev freelance part-time)', '—', '800 – 1,200'],
    ['Marketing inicial (Facebook + Google Ads)', '1,000', '500 – 1,000'],
    ['TOTALES', 'S/ 21,120 – 33,120', 'S/ 2,660 – 4,260'],
]
add_table(cost_headers, cost_rows)

p = doc.add_paragraph()
run = p.add_run('Costo operativo promedio mensual estimado: S/ 3,500/mes')
run.bold = True

add_orange_heading('Capacidad', 14)
cap_headers = ['Recurso', 'Límite']
cap_rows = [
    ['Restaurantes (negocios)', '50 – 80'],
    ['Usuarios concurrentes', '100 – 150'],
    ['Pedidos por hora', '200 – 300'],
    ['Respuestas IA por minuto', '10 – 15'],
    ['Archivos almacenados', '~5,000'],
    ['Artículos de inventario por negocio', '500+'],
]
add_table(cap_headers, cap_rows)

add_orange_heading('Rentabilidad', 14)
add_body('Planes de precios:')
plans = [
    ('Free', '0', '1 mesa, 10 pedidos/día — Prueba'),
    ('Básico', '49', 'Hasta 10 mesas, inventario básico'),
    ('Pro', '99', 'Hasta 30 mesas, todo incluido'),
    ('Enterprise', '249', 'Ilimitado, multi-sede'),
]
add_table(['Plan', 'Precio (S//mes)', 'Descripción'], plans)

add_body('Precio promedio ponderado estimado: S/ 75 / cliente / mes')

roe_headers = ['Clientes', 'Ingreso mensual', 'Ganancia / Pérdida']
roe_rows = [
    ['10', 'S/ 750', '– S/ 2,750'],
    ['20', 'S/ 1,500', '– S/ 2,000'],
    ['30', 'S/ 2,250', '– S/ 1,250'],
    ['40', 'S/ 3,000', '– S/ 500'],
    ['47 (BREAK-EVEN)', 'S/ 3,525', '≈ S/ 0'],
    ['50', 'S/ 4,375', '+ S/ 875'],
    ['60', 'S/ 4,500', '+ S/ 1,000'],
    ['80', 'S/ 6,000', '+ S/ 2,500'],
]
add_table(roe_headers, roe_rows)

p = doc.add_paragraph()
run = p.add_run('Break-even: ~47 clientes (o ~36 en plan Pro)')
run.bold = True
p.add_run('\nRecuperación de inversión inicial (S/ 27k): ~18 meses con 60 clientes')

add_body('Tiempo de desarrollo: 6 – 10 semanas')

doc.add_page_break()

# --- SECTION 2: VERSIÓN COMPLETA ---
add_orange_heading('VERSIÓN 2: COMPLETA (Web + App Mobile + Cloud Escalable)', 18)
add_body('Plataforma completa con app nativa iOS/Android (Flutter/React Native), infraestructura cloud escalable, seguridad profesional. Soporta cientos de restaurantes sin lentitud.')

add_orange_heading('Funcionalidades incluidas (adicionales a la básica)', 14)
features_comp = [
    'App nativa para clientes (pedir desde el celular con notificaciones push)',
    'App nativa para empresa (gestión desde el móvil sin abrir navegador)',
    'Dashboard avanzado con reportes exportables (PDF/Excel) y gráficos',
    'Infraestructura cloud escalable (auto-scaling, balanceo de carga)',
    'Integración directa OSE Sunat + backup',
    'Yape multilink como método de pago adicional',
    'IA Chat con fine-tuning (respuestas más precisas por restaurante)',
    'Seguridad profesional: WAF, pentesting, auditoría OWASP, rate limiting',
]
for f in features_comp:
    add_bullet(f)

add_orange_heading('Infraestructura', 14)
infra_comp = [
    'AWS / Google Cloud (multi-AZ)',
    'RDS / Cloud SQL (base de datos gerenciada)',
    'CloudFront + S3 (CDN + almacenamiento)',
    'Auto-scaling según demanda',
    'Redis para WebSockets y caché',
]
for i in infra_comp:
    add_bullet(i)

add_orange_heading('Costos', 14)
cost_rows2 = [
    ['Backend escalable (FastAPI + Redis + Workers + CDN)', '15,000 – 22,000', '—'],
    ['Frontend web (Next.js + SSR + PWA avanzada)', '8,000 – 14,000', '—'],
    ['App mobile nativa (Flutter, iOS + Android)', '16,000 – 28,000', '—'],
    ['Facturación Sunat (integración directa OSE + backup)', '3,000 – 4,500', '—'],
    ['Pasarela de pagos (MP + Culqi + Yape multilink)', '2,500 – 4,000', '—'],
    ['Multi-sede + jerarquía + roles + permisos (full)', '2,500 – 4,000', '—'],
    ['IA Chat (GPT-4o-mini fine-tuned)', '—', '1,500 – 2,500'],
    ['Infraestructura cloud (AWS/GCP escalable)', '—', '600 – 1,200'],
    ['Base de datos gerenciada (RDS/Cloud SQL + backups)', '—', '300 – 600'],
    ['CDN + Storage (CloudFront + S3)', '—', '100 – 200'],
    ['Seguridad (WAF, pentesting, OWASP audit)', '4,000 – 8,000', '200 – 400'],
    ['Dominio .pe + SSL + emails corporativos', '200', '40'],
    ['UX/UI premium (diseño profesional)', '2,000 – 4,000', '—'],
    ['Landing page profesional + SEO + blog + demo', '2,000 – 4,000', '—'],
    ['Dashboard avanzado (reportes exportables)', '1,500 – 3,000', '—'],
    ['Soporte técnico (1 persona full-time L1/L2)', '—', '2,000 – 3,000'],
    ['Marketing (Ads + Google Play + ASO + TikTok)', '3,000', '2,000 – 4,000'],
    ['TOTALES', 'S/ 59,700 – 97,700', 'S/ 6,700 – 11,900'],
]
add_table(cost_headers, cost_rows2)

p = doc.add_paragraph()
run = p.add_run('Costo operativo promedio mensual estimado: S/ 9,300/mes')
run.bold = True

add_orange_heading('Capacidad', 14)
cap_rows2 = [
    ['Restaurantes (negocios)', '500+ (escalamiento horizontal)'],
    ['Usuarios concurrentes', '2,000+'],
    ['Pedidos por hora', '5,000+'],
    ['Respuestas IA por minuto', '100+ (fine-tuned)'],
    ['Archivos almacenados', 'Ilimitado (S3)'],
    ['Facturación por minuto', '50+'],
]
add_table(cap_headers, cap_rows2)

add_orange_heading('Rentabilidad', 14)
roe_rows2 = [
    ['20', 'S/ 1,500', '– S/ 7,800'],
    ['50', 'S/ 3,750', '– S/ 5,550'],
    ['80', 'S/ 6,000', '– S/ 3,300'],
    ['124 (BREAK-EVEN)', 'S/ 9,300', '≈ S/ 0'],
    ['150', 'S/ 11,250', '+ S/ 1,950'],
    ['200', 'S/ 15,000', '+ S/ 5,700'],
    ['500', 'S/ 37,500', '+ S/ 28,200'],
]
add_table(roe_headers, roe_rows2)

p = doc.add_paragraph()
run = p.add_run('Break-even: ~124 clientes (o ~94 en plan Pro)')
run.bold = True
p.add_run('\nRecuperación de inversión inicial (S/ 78k): ~16 meses con 200 clientes')

add_body('Tiempo de desarrollo: 4 – 6 meses')

doc.add_page_break()

# --- SECTION 3: CUADRO COMPARATIVO ---
add_orange_heading('CUADRO COMPARATIVO', 18)

comp_headers = ['Aspecto', 'Versión Básica', 'Versión Completa']
comp_rows = [
    ['Plataforma', 'Web PWA (navegador)', 'Web PWA + Apps nativas iOS/Android'],
    ['Instalación', 'Sin tienda (solo URL)', 'Google Play + App Store'],
    ['Notificaciones push', 'Solo web (navegador)', 'Push nativas en móvil'],
    ['Facturación Sunat', '✅ vía API (Nubefact)', '✅ OSE directa + backup'],
    ['Pagos', '✅ MP + Culqi', '✅ MP + Culqi + Yape multilink'],
    ['Multi-sede + roles + permisos', '✅ Completo', '✅ Completo'],
    ['IA Chat', '✅ GPT-4o-mini', '✅ GPT-4o-mini fine-tuned'],
    ['Seguridad', 'SSL + básica', 'WAF + pentesting + OWASP'],
    ['Escalabilidad', 'Limitada (1 VPS)', 'Ilimitada (cloud auto-scaling)'],
    ['Soporte', 'Dev part-time', '1 full-time + SLA'],
    ['Dashboard', 'Estándar', 'Avanzado + reportes exportables'],
    ['UX/UI', 'Plantilla funcional', 'Diseño premium'],
    ['Inversión inicial', 'S/ 21k – 33k', 'S/ 60k – 98k'],
    ['Costo operativo mensual', 'S/ 2.7k – 4.3k', 'S/ 6.7k – 11.9k'],
    ['Costo promedio mensual', '~ S/ 3,500', '~ S/ 9,300'],
    ['Capacidad máxima', '50 – 80 restaurantes', '500+ restaurantes'],
    ['Break-even (S/ 75 prom.)', '~47 clientes', '~124 clientes'],
    ['Break-even (solo Pro S/ 99)', '~36 clientes', '~94 clientes'],
    ['Recuperación inversión', '~18 meses (60 clientes)', '~16 meses (200 clientes)'],
    ['Tiempo desarrollo', '6 – 10 semanas', '4 – 6 meses'],
]
add_table(comp_headers, comp_rows)

# --- SECTION 4: ESTRATEGIA ---
add_orange_heading('ESTRATEGIA RECOMENDADA EN 3 FASES', 18)

add_orange_heading('Fase 1 — Básico (Mes 1 – 3)', 14)
f1 = [
    'Invertir S/ 27k en desarrollo del MVP completo',
    '10 clientes beta gratis para validación',
    'Costo mensual: S/ 3,500',
    'Objetivo: 15 – 20 clientes pagos',
]
for f in f1:
    add_bullet(f)

add_orange_heading('Fase 2 — Crecimiento (Mes 4 – 8)', 14)
f2 = [
    'Escalar a 40 – 50 clientes con marketing',
    'Al llegar a 50 – 60 clientes, evaluar necesidad de app nativa',
    'Ingreso mensual: S/ 3,000 – 4,500',
    'Break-even alcanzado (~47 clientes)',
]
for f in f2:
    add_bullet(f)

add_orange_heading('Fase 3 — Salto a Completo (Mes 9 – 12)', 14)
f3 = [
    'Con flujo de caja de 50+ clientes (~S/ 4,500+/mes)',
    'Financiar el desarrollo mobile y migración a cloud',
    'Inversión adicional: S/ 40k – 70k',
    'Objetivo: 150+ clientes en 12 meses',
]
for f in f3:
    add_bullet(f)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('QuickTable — Menú digital inteligente para restaurantes')
run.font.italic = True
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run.font.size = Pt(10)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Junio 2026')
run2.font.italic = True
run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
run2.font.size = Pt(10)

doc.add_page_break()

# ============================================================
# SECTION 4: PLAN DE DISTRIBUCIÓN
# ============================================================
add_orange_heading('PLAN DE DISTRIBUCIÓN', 18)

add_orange_heading('Canales de adquisición de clientes', 14)
add_table(
    ['Canal', 'Costo inicial (S/)', 'Costo recurrente (S//mes)', 'Alcance estimado', 'Tasa de conversión'],
    [
        ['Venta directa (visitando restaurantes)', '0', '2,500 – 4,000', '30 – 50 visitas/mes', '15 – 25%'],
        ['Facebook / Instagram Ads', '500', '1,000 – 2,000', '50k – 150k imp./mes', '2 – 5%'],
        ['Google Ads', '500', '1,000 – 2,000', '20 – 50 clics/día', '5 – 10%'],
        ['TikTok orgánico', '0', '0', 'Viral potencial', 'Alto'],
        ['Aliados estratégicos', '300', '200 – 500', '10 – 30 leads/mes', '10 – 20%'],
        ['Marketplace (Google Play)', '0', '0 (comisión 15%)', 'Tráfico orgánico', 'Variable'],
        ['Referidos (10% dto.)', '0', '10% del MRR', 'Boca a boca', '20 – 30%'],
    ]
)

add_orange_heading('Estrategia por fase', 14)

add_orange_heading('Mes 1 – 3: Siembra', 12)
f_dist_1 = [
    'Venta directa intensiva en Lima: Miraflores, Barranco, San Isidro, Surco, San Borja',
    '1 vendedor visitando 8 – 10 restaurantes por día',
    'Objetivo: 15 – 20 clientes beta gratis → primeros 10 pagos',
    'Aliados: contactar 3 distribuidoras de insumos',
    'TikTok: 4 – 6 videos/semana (tutoriales, tips de inventario, cocina en acción)',
]
for f in f_dist_1:
    add_bullet(f)

add_orange_heading('Mes 4 – 8: Crecimiento', 12)
f_dist_2 = [
    'Google Ads + Facebook Ads activos con S/ 1,500 – 2,000/mes cada uno',
    '2 vendedores en campo (Lima + Callao)',
    'Primeros casos de éxito → contenido para ads y TikTok',
    'Programa de referidos activo',
    'Objetivo: 40 – 60 clientes pagos',
]
for f in f_dist_2:
    add_bullet(f)

add_orange_heading('Mes 9 – 12: Expansión', 12)
f_dist_3 = [
    'Expandir venta directa a provincias: Arequipa, Trujillo, Cusco, Chiclayo (Zoom + viajes quincenales)',
    'Publicidad en Facebook segmentada por ciudad',
    'Aliados estratégicos nacionales (Cámaras de Comercio regionales)',
    'Google Play Store si se lanza app nativa',
    'Objetivo: 100 – 150 clientes pagos',
]
for f in f_dist_3:
    add_bullet(f)

add_orange_heading('Presupuesto de distribución (acumulado 12 meses)', 14)
add_table(
    ['Rubro', 'Mes 1-3', 'Mes 4-8', 'Mes 9-12', 'Total 12 meses'],
    [
        ['Vendedores', 'S/ 7,500', 'S/ 20,000', 'S/ 25,000', 'S/ 52,500'],
        ['Facebook/Instagram Ads', 'S/ 1,500', 'S/ 7,500', 'S/ 9,000', 'S/ 18,000'],
        ['Google Ads', 'S/ 1,500', 'S/ 7,500', 'S/ 9,000', 'S/ 18,000'],
        ['Aliados estratégicos', 'S/ 600', 'S/ 2,000', 'S/ 2,000', 'S/ 4,600'],
        ['TikTok (producción)', 'S/ 500', 'S/ 1,500', 'S/ 1,500', 'S/ 3,500'],
        ['Total distribución', 'S/ 11,600', 'S/ 38,500', 'S/ 46,500', 'S/ 96,600'],
    ]
)

p_note = doc.add_paragraph()
run_note = p_note.add_run('Nota: Estos costos están incluidos en los rubros de marketing y soporte de los costos generales.')
run_note.font.italic = True
run_note.font.size = Pt(10)
run_note.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# SECTION 5: TIEMPO DE RECUPERACIÓN DE INVERSIÓN
# ============================================================
add_orange_heading('TIEMPO DE RECUPERACIÓN DE INVERSIÓN (ROI)', 18)

add_orange_heading('Escenario Básico (Inversión: S/ 27,000 | Costo mensual: S/ 3,500)', 14)

roi_headers = ['Mes', 'Clientes', 'Ingreso mensual', 'Costo mensual', 'Flujo del mes', 'Flujo acum.', 'Inversión rest.']
roi_basic_rows = [
    ['1', '0 (desarrollo)', '0', '3,500', '–3,500', '–30,500', '30,500'],
    ['2', '0 (desarrollo)', '0', '3,500', '–3,500', '–34,000', '34,000'],
    ['3', '10 (beta)', '0', '3,500', '–3,500', '–37,500', '37,500'],
    ['4', '15', '1,125', '3,500', '–2,375', '–39,875', '39,875'],
    ['5', '22', '1,650', '3,500', '–1,850', '–41,725', '41,725'],
    ['6', '30', '2,250', '3,500', '–1,250', '–42,975', '42,975'],
    ['7', '38', '2,850', '3,500', '–650', '–43,625', '43,625'],
    ['8', '47', '3,525', '3,500', '+25', '–43,600', '43,600'],
    ['9', '55', '4,125', '3,500', '+625', '–42,975', '42,975'],
    ['10', '62', '4,650', '3,500', '+1,150', '–41,825', '41,825'],
    ['11', '70', '5,250', '3,500', '+1,750', '–40,075', '40,075'],
    ['12', '80', '6,000', '3,500', '+2,500', '–37,575', '37,575'],
    ['13', '85', '6,375', '3,700', '+2,675', '–34,900', '34,900'],
    ['14', '90', '6,750', '3,700', '+3,050', '–31,850', '31,850'],
    ['15', '95', '7,125', '3,700', '+3,425', '–28,425', '28,425'],
    ['16', '100', '7,500', '3,700', '+3,800', '–24,625', '24,625'],
    ['17', '105', '7,875', '3,700', '+4,175', '–20,450', '20,450'],
    ['18', '110', '8,250', '3,700', '+4,550', '–15,900', '15,900'],
    ['19', '115', '8,625', '3,700', '+4,925', '–10,975', '10,975'],
    ['20', '120', '9,000', '3,700', '+5,300', '–5,675', '5,675'],
    ['21 ✅', '125', '9,375', '3,700', '+5,675', '≈ 0', '✅ Recuperado'],
    ['22', '130', '9,750', '3,800', '+5,950', '+5,950', 'Ganancia'],
]
add_table(roi_headers, roi_basic_rows)

p = doc.add_paragraph()
run = p.add_run('Recuperación de inversión (Escenario Básico): 21 meses (1 año 9 meses)')
run.bold = True
run.font.size = Pt(12)

add_body('Asumiendo: precio promedio S/ 75/cliente, crecimiento de 5-10 clientes nuevos/mes.')

add_orange_heading('Escenario Básico — Sensibilidad por precio', 12)
add_table(
    ['Precio promedio', 'Clientes necesarios', 'Tiempo recuperación'],
    [
        ['S/ 49 (solo básico)', '72', '~36 meses'],
        ['S/ 75 (mix de planes)', '47', '~21 meses'],
        ['S/ 99 (todos Pro)', '36', '~16 meses'],
        ['S/ 129 (Pro + addons)', '28', '~13 meses'],
    ]
)

doc.add_page_break()

add_orange_heading('Escenario Completo (Inversión: S/ 78,000 | Costo mensual: S/ 9,300)', 14)

roi_comp_rows = [
    ['1-4', '0 (desarrollo)', '0', '9,300', '–37,200', '–115,200', '115,200'],
    ['5', '5 (beta)', '0', '9,300', '–9,300', '–124,500', '124,500'],
    ['6', '15', '1,125', '9,300', '–8,175', '–132,675', '132,675'],
    ['7', '25', '1,875', '9,300', '–7,425', '–140,100', '140,100'],
    ['8', '38', '2,850', '9,300', '–6,450', '–146,550', '146,550'],
    ['9', '50', '3,750', '9,300', '–5,550', '–152,100', '152,100'],
    ['10', '65', '4,875', '9,300', '–4,425', '–156,525', '156,525'],
    ['11', '80', '6,000', '9,300', '–3,300', '–159,825', '159,825'],
    ['12', '95', '7,125', '9,500', '–2,375', '–162,200', '162,200'],
    ['13', '110', '8,250', '9,500', '–1,250', '–163,450', '163,450'],
    ['14', '124', '9,300', '9,500', '–200', '–163,650', '163,650'],
    ['15', '140', '10,500', '9,500', '+1,000', '–162,650', '162,650'],
    ['16', '155', '11,625', '9,500', '+2,125', '–160,525', '160,525'],
    ['17', '170', '12,750', '9,500', '+3,250', '–157,275', '157,275'],
    ['18', '185', '13,875', '9,500', '+4,375', '–152,900', '152,900'],
    ['19', '200', '15,000', '9,500', '+5,500', '–147,400', '147,400'],
    ['20', '215', '16,125', '9,500', '+6,625', '–140,775', '140,775'],
    ['21', '230', '17,250', '9,800', '+7,450', '–133,325', '133,325'],
    ['22', '245', '18,375', '9,800', '+8,575', '–124,750', '124,750'],
    ['23', '260', '19,500', '9,800', '+9,700', '–115,050', '115,050'],
    ['24', '275', '20,625', '9,800', '+10,825', '–104,225', '104,225'],
    ['25', '290', '21,750', '10,000', '+11,750', '–92,475', '92,475'],
    ['26', '305', '22,875', '10,000', '+12,875', '–79,600', '79,600'],
    ['27', '320', '24,000', '10,000', '+14,000', '–65,600', '65,600'],
    ['28', '335', '25,125', '10,000', '+15,125', '–50,475', '50,475'],
    ['29', '350', '26,250', '10,000', '+16,250', '–34,225', '34,225'],
    ['30', '365', '27,375', '10,000', '+17,375', '–16,850', '16,850'],
    ['31 ✅', '380', '28,500', '10,000', '+18,500', '≈ 0', '✅ Recuperado'],
]
add_table(roi_headers, roi_comp_rows)

p = doc.add_paragraph()
run = p.add_run('Recuperación de inversión (Escenario Completo): 31 meses (2 años 7 meses)')
run.bold = True
run.font.size = Pt(12)

add_body('Nota: Requiere 4 meses de desarrollo sin ingresos. Una vez superado el break-even, las ganancias escalan más rápido.')

doc.add_page_break()

# ============================================================
# SECTION 6: PRECIO ÓPTIMO PARA EL MERCADO PERUANO
# ============================================================
add_orange_heading('PRECIO ÓPTIMO PARA EL MERCADO PERUANO', 18)

add_orange_heading('Análisis de competencia directa en Perú', 14)
add_table(
    ['Competidor', 'Tipo', 'Precio (US$/mes)', 'Precio (S//mes)', 'Ideal para'],
    [
        ['Venko', 'ERP gastronómico completo', '50 – 150', '185 – 555', 'Rest. grandes'],
        ['WasiPOS', 'POS + inventario', '40 – 80', '148 – 296', 'Pequeños/medianos'],
        ['Zenda', 'Delivery + POS', '30 – 70', '111 – 259', 'Delivery-first'],
        ['Urbaner POS', 'POS + cocina', '25 – 60', '93 – 222', 'Fast food'],
        ['TuCarta.pe', 'Solo menú digital QR', '8 – 25', '30 – 93', 'Muy pequeños'],
        ['Sisfoh', 'ERP general', '20 – 50', '74 – 185', 'Cualquier rubro'],
        ['QuickTable (propuesto)', 'Todo en uno', '—', '49 – 249', 'Pequeños / medianos'],
    ]
)

add_orange_heading('Elasticidad del mercado por segmento', 14)

add_orange_heading('Restaurante pequeño (1 – 5 mesas)', 12)
f_small = [
    'Presupuesto para software: S/ 0 – 50/mes',
    'Dolor: "no tengo tiempo para controles"',
    'Solo paga si ve ROI claro en ahorro de tiempo',
    'Precio máximo: S/ 49/mes (plan Básico)',
]
for f in f_small:
    add_bullet(f)

add_orange_heading('Restaurante mediano (6 – 20 mesas) — segmento TARGET', 12)
f_med = [
    'Presupuesto para software: S/ 50 – 150/mes',
    'Dolor: "pierdo insumos, no sé cuánto gano, cocina se atrasa"',
    'Ya usan Excel o POS básico',
    'Precio óptimo: S/ 99/mes (plan Pro)',
]
for f in f_med:
    add_bullet(f)

add_orange_heading('Restaurante grande o cadena (20+ mesas)', 12)
f_large = [
    'Presupuesto para software: S/ 150 – 500/mes',
    'Dolor: "necesito control centralizado, facturación, reportes"',
    'Ya evalúan Venko o WasiPOS',
    'Precio óptimo: S/ 249/mes (plan Enterprise)',
]
for f in f_large:
    add_bullet(f)

add_orange_heading('Precio recomendado', 14)
add_table(
    ['Plan', 'Precio (S//mes)', 'Posicionamiento'],
    [
        ['Free', 'S/ 0', 'Gancho de entrada — 1 mesa, 10 pedidos/día'],
        ['Básico', 'S/ 49', 'Restaurante pequeño que recién empieza'],
        ['Pro 🔥', 'S/ 99', 'Caballo de batalla — mejor relación valor/precio'],
        ['Enterprise', 'S/ 249', 'Cadenas y restaurantes grandes'],
    ]
)

p = doc.add_paragraph()
run = p.add_run('Precio óptimo recomendado: S/ 99/mes (plan Pro)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xC2, 0x41, 0x0C)

add_orange_heading('¿Por qué S/ 99?', 12)
why_reasons = [
    'Por debajo de WasiPOS (S/ 148) y Venko (S/ 185) — competimos en precio',
    'Por encima de TuCarta.pe (S/ 30-93) — ofrecemos mucho más valor',
    'Precio psicológico: S/ 99 suena a "menos de 100 soles" — umbral clave en Perú',
    'Margen saludable: 50 clientes Pro generan S/ 4,950/mes contra S/ 3,500 de costos = 40% margen',
    'Comparable a: una cena para dos en restaurante de gama media — los dueños lo entienden como "lo que gano en un día"',
]
for r in why_reasons:
    add_bullet(r)

add_orange_heading('Estrategia de precios a 12 meses', 14)
add_table(
    ['Mes', 'Precio Pro', 'Estrategia'],
    [
        ['1 – 6', 'S/ 69 (lanzamiento)', 'Penetración agresiva para ganar tracción'],
        ['7 – 12', 'S/ 99 (regular)', 'Subida gradual con casos de éxito'],
        ['12+', 'S/ 119 – 129', 'Consolidación con mejoras y marca'],
    ]
)

add_orange_heading('Resumen de precio óptimo', 14)
add_table(
    ['Variable', 'Valor'],
    [
        ['Precio plan Pro', 'S/ 99/mes'],
        ['Precio de lanzamiento', 'S/ 69/mes (primeros 6 meses)'],
        ['Competidor directo', 'WasiPOS (S/ 148/mes) + TuCarta.pe (S/ 30-93/mes)'],
        ['Ventaja diferencial', 'Todo en uno: menú + cocina + inventario + facturación + IA'],
        ['Segmento target', 'Restaurantes medianos (6-20 mesas) en Lima y principales ciudades'],
        ['Valor percibido', 'S/ 99 = ~1 cubierto diario → "si me ahorra un plato robado al mes, ya pagó"'],
    ]
)

# --- Save ---
output_path = '/data/.openclaw/workspace/quicktable/PLAN-COMERCIALIZACION.docx'
doc.save(output_path)
print(f'✅ Documento guardado: {output_path}')
